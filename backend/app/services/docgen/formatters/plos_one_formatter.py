"""
app/services/docgen/formatters/plos_one_formatter.py

PLOS ONE journal formatter.

Reads ALL formatting values from the template's JSONB columns.

PLOS ONE format profile (as seeded in template_seeder.py):
  - Margins: 2.54 cm all sides
  - Font: Arial 11pt, double-spaced
  - Max abstract: 300 words (unstructured single paragraph, no citations)
  - Heading levels: 1, 2, and 3
  - Citation style: bracketed_numbers [1], [2]
  - Required statements: data_availability (mandatory), ethics_statement
  - No abstract citations (validated at preflight, flagged here if detected)
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.manuscript_ir import ManuscriptIR
from .base_formatter import BaseFormatter

if TYPE_CHECKING:
    from app.services.docgen.reference_formatter import ReferenceFormatter


_DEFAULTS = {
    "font_name": "Arial",
    "font_size_pt": 11.0,
    "line_spacing_pt": 22.0,   # double-spaced at 11pt
    "margin_cm": 2.54,
    "space_before_pt": 0.0,
    "space_after_pt": 6.0,
}


class PlosOneFormatter(BaseFormatter):
    """Formatter for the PLOS ONE journal template (slug: 'plos-one')."""

    # ── Styling ────────────────────────────────────────────────────────────────

    def apply_document_style(self, doc: Document, rules: dict) -> None:
        margin = rules.get("margin_cm", _DEFAULTS["margin_cm"])
        font_name = rules.get("font_name", _DEFAULTS["font_name"])
        font_size = rules.get("font_size_pt", _DEFAULTS["font_size_pt"])
        line_spacing = rules.get("line_spacing_pt", _DEFAULTS["line_spacing_pt"])

        self._set_margins(doc, margin, margin, margin, margin)
        self._set_default_font(doc, font_name, font_size)
        self._set_paragraph_spacing(
            doc,
            space_before_pt=_DEFAULTS["space_before_pt"],
            space_after_pt=rules.get("space_after_pt", _DEFAULTS["space_after_pt"]),
            line_spacing_pt=line_spacing,
        )

    # ── Title Page ─────────────────────────────────────────────────────────────

    def build_title_page(self, doc: Document, ir: ManuscriptIR, layout: dict) -> None:
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(ir.title)
        run.bold = True
        run.font.size = Pt(14)

        # Authors with superscript affiliation indices
        if ir.authors:
            authors_para = doc.add_paragraph()
            authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, author in enumerate(ir.authors):
                if i > 0:
                    authors_para.add_run(", ")
                name_run = authors_para.add_run(f"{author.given_name} {author.surname}")
                if author.affiliation_indices:
                    idx_run = authors_para.add_run(",".join(str(j) for j in author.affiliation_indices))
                    idx_run.font.superscript = True

        # Affiliations
        for aff in ir.affiliations:
            idx_run_para = doc.add_paragraph()
            idx_run_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            superscript_run = idx_run_para.add_run(str(aff.index))
            superscript_run.font.superscript = True
            parts = [aff.institution]
            if aff.department:
                parts.append(aff.department)
            if aff.city:
                parts.append(aff.city)
            if aff.country:
                parts.append(aff.country)
            idx_run_para.add_run(" " + ", ".join(parts))

        # Corresponding author
        if ir.corresponding_author:
            ca = ir.corresponding_author
            corr_template = layout.get("corresponding_author", "Corresponding author")
            doc.add_paragraph(f"{corr_template}: {ca.full_name}, {ca.email}")

        self._add_page_break(doc)

    # ── Abstract ───────────────────────────────────────────────────────────────

    def build_abstract(self, doc: Document, ir: ManuscriptIR) -> None:
        doc.add_heading("Abstract", level=1)

        import re
        has_citations = bool(re.search(r'\[\d+\]|\(\w+\s+et\s+al', ir.abstract or ""))
        abstract_text = ir.abstract
        if has_citations:
            abstract_text = "[PREFLIGHT WARNING: Abstract may contain citations — PLOS ONE forbids citations in the abstract]\n" + abstract_text

        doc.add_paragraph(abstract_text)

        if ir.keywords:
            kw_para = doc.add_paragraph()
            kw_para.add_run("Keywords: ").bold = True
            kw_para.add_run(", ".join(ir.keywords))

    # ── References ────────────────────────────────────────────────────────────

    def build_references(
        self,
        doc: Document,
        refs: list,
        ref_formatter: "ReferenceFormatter",
    ) -> None:
        if not refs:
            return
        doc.add_heading("References", level=1)
        for ref in sorted(refs, key=lambda r: r.index):
            entry = ref_formatter.format_entry(ref)
            doc.add_paragraph(entry, style="List Number")

    # ── Required Statements ───────────────────────────────────────────────────

    def build_statements(self, doc: Document, ir: ManuscriptIR, required: dict) -> None:
        has_any = any([
            ir.data_availability,
            ir.ethics_statement,
            ir.conflict_of_interest,
            ir.acknowledgements,
            ir.funding,
        ])
        if not has_any:
            return

        doc.add_heading("Supporting Information", level=1)

        # Data Availability is mandatory for PLOS ONE
        if ir.data_availability:
            self._add_labeled_statement(doc, "Data Availability Statement", ir.data_availability)
        else:
            doc.add_paragraph(
                "[MISSING — PLOS ONE requires a Data Availability Statement for every submission]"
            )

        self._add_labeled_statement(doc, "Ethics Statement", ir.ethics_statement)
        self._add_labeled_statement(doc, "Competing Interests", ir.conflict_of_interest)
        self._add_labeled_statement(doc, "Acknowledgements", ir.acknowledgements)

        if ir.funding:
            funding_parts = []
            for fs in ir.funding:
                part = fs.funder
                if fs.grant_number:
                    part += f" (grant {fs.grant_number})"
                funding_parts.append(part)
            self._add_labeled_statement(doc, "Funding", "; ".join(funding_parts))
