"""
app/services/docgen/formatters/base_formatter.py

Abstract Base Class for journal-specific .docx formatters.

Every concrete formatter receives fully-typed ManuscriptIR data and reads
formatting parameters directly from the JournalTemplate's JSON config columns
(formatting_rules, title_page_layout, reference_format, required_statements).
No journal-specific values are hardcoded in Python — the DB row is the
source of truth.
"""

from __future__ import annotations

from abc import ABC, abstractmethod
from io import BytesIO
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm

from app.schemas.manuscript_ir import ManuscriptIR, SectionNode

if TYPE_CHECKING:
    from app.services.docgen.reference_formatter import ReferenceFormatter


class BaseFormatter(ABC):
    """
    Abstract base for all journal formatters.

    Concrete subclasses implement the abstract methods (apply_document_style,
    build_title_page, build_abstract, build_references, build_statements).
    The recursive build_sections implementation is provided here and may be
    overridden when a journal requires non-standard section rendering.
    """

    # ── Helpers ────────────────────────────────────────────────────────────────

    @staticmethod
    def _set_margins(doc: Document, top_cm: float, bottom_cm: float, left_cm: float, right_cm: float) -> None:
        """Apply page margins to all sections in the document."""
        for section in doc.sections:
            section.top_margin    = Cm(top_cm)
            section.bottom_margin = Cm(bottom_cm)
            section.left_margin   = Cm(left_cm)
            section.right_margin  = Cm(right_cm)

    @staticmethod
    def _set_default_font(doc: Document, font_name: str, font_size_pt: float) -> None:
        """Set default body font for the document normal style."""
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size_pt)

    @staticmethod
    def _set_paragraph_spacing(doc: Document, space_before_pt: float, space_after_pt: float, line_spacing_pt: float | None = None) -> None:
        """Set paragraph spacing on the Normal style."""
        style = doc.styles["Normal"]
        pf = style.paragraph_format
        pf.space_before = Pt(space_before_pt)
        pf.space_after  = Pt(space_after_pt)
        if line_spacing_pt is not None:
            from docx.shared import Pt as _Pt
            pf.line_spacing = _Pt(line_spacing_pt)

    @staticmethod
    def _add_page_break(doc: Document) -> None:
        doc.add_page_break()

    @staticmethod
    def _add_section_heading(doc: Document, text: str, level: int) -> None:
        """Add a heading, clamping to levels 1-4 (python-docx limit)."""
        doc.add_heading(text, level=max(1, min(4, level if level > 0 else 1)))

    @staticmethod
    def _add_labeled_statement(doc: Document, label: str, content: str | None) -> None:
        """Add a labeled required-statement paragraph (e.g. 'Conflict of Interest:')."""
        if not content:
            return
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(content)

    # ── Abstract Interface ─────────────────────────────────────────────────────

    @abstractmethod
    def apply_document_style(self, doc: Document, rules: dict) -> None:
        """
        Apply global document styling: page margins, base font, line spacing.
        Must read parameters entirely from `rules` (the template's
        formatting_rules JSONB column) — no hardcoded values.
        """

    @abstractmethod
    def build_title_page(self, doc: Document, ir: ManuscriptIR, layout: dict) -> None:
        """
        Build the title page (title, authors, affiliations, corresponding
        author). Read display rules from `layout` (title_page_layout column).
        """

    @abstractmethod
    def build_abstract(self, doc: Document, ir: ManuscriptIR) -> None:
        """Build the abstract section with keywords."""

    def build_sections(self, doc: Document, sections: list[SectionNode]) -> None:
        """
        Default recursive section builder.

        Iterates the SectionNode tree depth-first, emitting headings and
        paragraph content. Level-0 nodes (preamble) are rendered without a
        heading. Override in concrete formatters for journal-specific behavior
        (e.g. two-column paragraphs, inline figure labels, LaTeX flag).

        KNOWN LIMITATION: LaTeX formula strings present in section.content
        are appended as raw text — no MathML / OMML rendering is attempted.
        They are flagged with a [FORMULA] prefix so reviewers can identify them.
        """
        for node in sections:
            if node.level == 0:
                # Preamble / root — no heading
                for para_text in node.content:
                    doc.add_paragraph(self._flag_formulas(para_text))
            else:
                self._add_section_heading(doc, node.heading, node.level)
                for para_text in node.content:
                    doc.add_paragraph(self._flag_formulas(para_text))
            if node.children:
                self.build_sections(doc, node.children)

    @abstractmethod
    def build_references(
        self,
        doc: Document,
        refs: list,
        ref_formatter: "ReferenceFormatter",
    ) -> None:
        """Build the references section using the configured citation style."""

    @abstractmethod
    def build_statements(self, doc: Document, ir: ManuscriptIR, required: dict) -> None:
        """Build mandatory disclosure statements (COI, ethics, data availability, etc.)."""

    # ── Formula Detection ─────────────────────────────────────────────────────

    @staticmethod
    def _flag_formulas(text: str) -> str:
        """
        Detect likely LaTeX fragments (\\frac, \\sum, $…$, $$…$$) and
        prefix them with [FORMULA] so they are visible rather than silently
        garbled. Full MathML rendering is a known Phase 2 limitation.
        """
        import re
        latex_pattern = re.compile(
            r'(\$\$.*?\$\$|\$[^$]+\$|\\[a-zA-Z]+\{|\\begin\{)',
            re.DOTALL
        )
        if latex_pattern.search(text):
            return f"[FORMULA — review rendering] {text}"
        return text
