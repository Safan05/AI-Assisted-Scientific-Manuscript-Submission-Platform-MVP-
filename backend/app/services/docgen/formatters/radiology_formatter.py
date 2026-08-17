"""
app/services/docgen/formatters/radiology_formatter.py

Radiology (RSNA) journal formatter.
"""
from __future__ import annotations

from typing import TYPE_CHECKING
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.manuscript_ir import ManuscriptIR
from .base_formatter import BaseFormatter

if TYPE_CHECKING:
    from app.services.docgen.reference_formatter import ReferenceFormatter

_DEFAULTS = {
    "font_name": "Times New Roman",
    "font_size_pt": 12.0,
    "line_spacing_pt": 24.0, # double-spaced
    "margin_cm": 2.54,
    "space_before_pt": 0.0,
    "space_after_pt": 0.0,
}

class RadiologyFormatter(BaseFormatter):
    """Formatter for the Radiology template (slug: 'radiology')."""

    def apply_document_style(self, doc: Document, rules: dict) -> None:
        margin = rules.get("margin_cm", _DEFAULTS["margin_cm"])
        font_name = rules.get("font_name", _DEFAULTS["font_name"])
        font_size = rules.get("font_size_pt", _DEFAULTS["font_size_pt"])
        line_spacing = rules.get("line_spacing_pt", _DEFAULTS["line_spacing_pt"])

        self._set_margins(doc, margin, margin, margin, margin)
        self._set_default_font(doc, font_name, font_size)
        self._set_paragraph_spacing(
            doc,
            space_before_pt=_DEFAULTS["space_before_pt"],
            space_after_pt=rules.get("space_after_pt", _DEFAULTS["space_after_pt"]),
            line_spacing_pt=line_spacing,
        )

    def build_title_page(self, doc: Document, ir: ManuscriptIR, layout: dict) -> None:
        # Title
        title_para = doc.add_paragraph()
        run = title_para.add_run(ir.title)
        run.bold = True
        
        # Authors
        author_names = [f"{a.given_name} {a.surname}" for a in ir.authors]
        doc.add_paragraph(", ".join(author_names))
        
        # Affiliations
        for aff in ir.affiliations:
            parts = [str(aff.index), aff.institution]
            if aff.department: parts.append(aff.department)
            if aff.city: parts.append(aff.city)
            if aff.country: parts.append(aff.country)
            doc.add_paragraph(". ".join(parts))

        # Correspondence
        if ir.corresponding_author:
            ca = ir.corresponding_author
            doc.add_paragraph(f"Address correspondence to: {ca.full_name} ({ca.email})")

        self._add_page_break(doc)

    def build_abstract(self, doc: Document, ir: ManuscriptIR) -> None:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(ir.abstract)
        self._add_page_break(doc)
