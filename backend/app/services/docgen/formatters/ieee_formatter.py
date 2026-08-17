"""
app/services/docgen/formatters/ieee_formatter.py

IEEE Transactions journal formatter.

IEEE format profile (as seeded in template_seeder.py):
  - Font: Times New Roman 10pt
  - Layout: Double column (approximated here by section properties if supported, 
    otherwise standard styling)
  - Headings: Roman numerals for top-level (I., II., III.)
  - Title page layout: Title, Authors, affiliations as footnotes (approximated)
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from app.schemas.manuscript_ir import ManuscriptIR
from .base_formatter import BaseFormatter

if TYPE_CHECKING:
    from app.services.docgen.reference_formatter import ReferenceFormatter


_DEFAULTS = {
    "font_name": "Times New Roman",
    "font_size_pt": 10.0,
    "line_spacing_pt": 12.0,   # single-spaced for IEEE double column
    "margin_cm": 1.9,          # tighter margins
    "space_before_pt": 6.0,
    "space_after_pt": 6.0,
}


def int_to_roman(num: int) -> str:
    """Convert integer to Roman numeral for IEEE headings."""
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
        ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
        ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num


class IEEEFormatter(BaseFormatter):
    """Formatter for the IEEE Transactions template (slug: 'ieee')."""

    def apply_document_style(self, doc: Document, rules: dict) -> None:
        margin = rules.get("margin_cm", _DEFAULTS["margin_cm"])
        font_name = rules.get("font_name", _DEFAULTS["font_name"])
        font_size = rules.get("font_size_pt", _DEFAULTS["font_size_pt"])
        line_spacing = rules.get("line_spacing_pt", _DEFAULTS["line_spacing_pt"])

        self._set_margins(doc, margin, margin, margin, margin)
        self._set_default_font(doc, font_name, font_size)
        self._set_paragraph_spacing(
            doc,
            space_before_pt=_DEFAULTS["space_before_pt"],
            space_after_pt=rules.get("space_after_pt", _DEFAULTS["space_after_pt"]),
            line_spacing_pt=line_spacing,
        )

    def build_title_page(self, doc: Document, ir: ManuscriptIR, layout: dict) -> None:
        # Title (Centered, larger font)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(ir.title)
        run.bold = True
        run.font.size = Pt(24)
        doc.add_paragraph() # spacer

        # Authors
        author_names = [f"{a.given_name} {a.surname}" for a in ir.authors]
        if author_names:
            authors_para = doc.add_paragraph(", ".join(author_names))
            authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Affiliations
        for aff in ir.affiliations:
            parts = [aff.institution]
            if aff.department: parts.append(aff.department)
            if aff.city: parts.append(aff.city)
            if aff.country: parts.append(aff.country)
            aff_para = doc.add_paragraph(", ".join(parts))
            aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in aff_para.runs:
                r.italic = True

        doc.add_paragraph() # spacer

    def build_abstract(self, doc: Document, ir: ManuscriptIR) -> None:
        abs_para = doc.add_paragraph()
        run = abs_para.add_run("Abstract—")
        run.bold = True
        run.italic = True
        abs_para.add_run(ir.abstract).italic = True

        if ir.keywords:
            kw_para = doc.add_paragraph()
            kw_para.add_run("Index Terms—").bold = True
            kw_para.add_run(", ".join(ir.keywords))

        doc.add_paragraph()

    def build_sections(self, doc: Document, sections: list) -> None:
        """Override to add Roman numerals to H1."""
        h1_counter = 1
        for section in sections:
            self._render_section(doc, section, 1, h1_counter)
            if section.level == 1:
                h1_counter += 1

    def _render_section(self, doc: Document, section, current_level: int, h1_counter: int) -> None:
        if section.heading:
            if current_level == 1:
                roman = int_to_roman(h1_counter)
                heading = f"{roman}. {section.heading.upper()}"
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(heading)
                run.bold = True
            elif current_level == 2:
                # IEEE H2: A. Heading Italic
                # We skip the A, B, C logic here for simplicity and just use italic
                para = doc.add_paragraph()
                run = para.add_run(section.heading)
                run.italic = True
            else:
                doc.add_heading(section.heading, level=current_level)

        for content_line in section.content:
            doc.add_paragraph(content_line)

        for child in section.children:
            self._render_section(doc, child, current_level + 1, h1_counter)

    def build_references(self, doc: Document, refs: list, ref_formatter: "ReferenceFormatter") -> None:
        if not refs:
            return
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("REFERENCES")
        run.bold = True

        for ref in sorted(refs, key=lambda r: r.index):
            entry = ref_formatter.format_entry(ref)
            p = doc.add_paragraph(f"[{ref.index}] {entry}")
            # Hanging indent for IEEE refs
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.first_line_indent = Pt(-18)

    def build_statements(self, doc: Document, ir: ManuscriptIR, required: dict) -> None:
        # IEEE usually puts acknowledgments at the end, right before references.
        # But this is called after references in generator.
        if ir.acknowledgements:
            doc.add_heading("ACKNOWLEDGMENT", level=1)
            doc.add_paragraph(ir.acknowledgements)
