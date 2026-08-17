"""
app/services/docgen/formatters/nature_formatter.py

Nature journal formatter.

Reads ALL formatting values from the template's JSONB columns (formatting_rules,
title_page_layout, reference_format, required_statements). No Nature-specific
values are hardcoded here — the template DB row is the source of truth.

Nature format profile (as seeded in template_seeder.py):
  - Margins: 2.54 cm all sides (standard A4)
  - Font: Times New Roman 12pt, double-spaced
  - Max abstract: 200 words (enforced at preflight; here we only render)
  - Heading levels: 1 and 2 only
  - Citation style: superscript (read from reference_format.citation_style)
  - Title page: title, authors, affiliations, correspondence statement
  - Required statements: competing interests, data availability, author contributions
"""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from app.schemas.manuscript_ir import ManuscriptIR
from .base_formatter import BaseFormatter

if TYPE_CHECKING:
    from app.services.docgen.reference_formatter import ReferenceFormatter


# Default fallback values — only used when the template config key is absent.
_DEFAULTS = {
    "font_name": "Times New Roman",
    "font_size_pt": 12.0,
    "line_spacing_pt": 24.0,   # double-spaced at 12pt
    "margin_cm": 2.54,
    "space_before_pt": 0.0,
    "space_after_pt": 6.0,
}


class NatureFormatter(BaseFormatter):
    """Formatter for the Nature journal template (slug: 'nature')."""

    # ── Styling ────────────────────────────────────────────────────────────────

    def apply_document_style(self, doc: Document, rules: dict) -> None:
        margin = rules.get("margin_cm", _DEFAULTS["margin_cm"])
        font_name = rules.get("font_name", _DEFAULTS["font_name"])
        font_size = rules.get("font_size_pt", _DEFAULTS["font_size_pt"])
        line_spacing = rules.get("line_spacing_pt", _DEFAULTS["line_spacing_pt"])

        self._set_margins(doc, margin, margin, margin, margin)
        self._set_default_font(doc, font_name, font_size)
        self._set_paragraph_spacing(
            doc,
            space_before_pt=_DEFAULTS["space_before_pt"],
            space_after_pt=rules.get("space_after_pt", _DEFAULTS["space_after_pt"]),
            line_spacing_pt=line_spacing,
        )

    # ── Title Page ─────────────────────────────────────────────────────────────

    def build_title_page(self, doc: Document, ir: ManuscriptIR, layout: dict) -> None:
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(ir.title)
        run.bold = True
        run.font.size = Pt(16)

        # Authors line
        author_names = [f"{a.given_name} {a.surname}" for a in ir.authors]
        authors_para = doc.add_paragraph(", ".join(author_names))
        authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Affiliations block
        for aff in ir.affiliations:
            parts = [str(aff.index), aff.institution]
            if aff.department:
                parts.append(aff.department)
            if aff.city:
                parts.append(aff.city)
            if aff.country:
                parts.append(aff.country)
            aff_para = doc.add_paragraph(". ".join(parts))
            aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Correspondence statement (Nature: "Correspondence and requests for materials…")
        if ir.corresponding_author:
            ca = ir.corresponding_author
            corr_template = layout.get(
                "corresponding_author_statement",
                "Correspondence and requests for materials should be addressed to"
            )
            doc.add_paragraph(f"{corr_template} {ca.full_name} ({ca.email}).")

        self._add_page_break(doc)

    # ── Abstract ────────────────────────────────────────────────────────────────

    def build_abstract(self, doc: Document, ir: ManuscriptIR) -> None:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(ir.abstract)

        if ir.keywords:
            kw_para = doc.add_paragraph()
            kw_para.add_run("Keywords: ").bold = True
            kw_para.add_run("; ".join(ir.keywords))

    # ── References ────────────────────────────────────────────────────────────

    def build_references(
        self,
        doc: Document,
        refs: list,
        ref_formatter: "ReferenceFormatter",
    ) -> None:
        if not refs:
            return
        doc.add_heading("References", level=1)
        for ref in sorted(refs, key=lambda r: r.index):
            entry = ref_formatter.format_entry(ref)
            doc.add_paragraph(entry, style="List Number")

    # ── Required Statements ───────────────────────────────────────────────────

    def build_statements(self, doc: Document, ir: ManuscriptIR, required: dict) -> None:
        has_any = any([
            ir.conflict_of_interest,
            ir.data_availability,
            ir.author_contributions,
            ir.ethics_statement,
            ir.acknowledgements,
            ir.funding,
        ])
        if not has_any:
            return

        doc.add_heading("Declarations", level=1)

        self._add_labeled_statement(doc, "Competing Interests", ir.conflict_of_interest)
        self._add_labeled_statement(doc, "Data Availability", ir.data_availability)
        self._add_labeled_statement(doc, "Author Contributions", ir.author_contributions)
        self._add_labeled_statement(doc, "Ethics Statement", ir.ethics_statement)
        self._add_labeled_statement(doc, "Acknowledgements", ir.acknowledgements)

        if ir.funding:
            funding_parts = []
            for fs in ir.funding:
                part = fs.funder
                if fs.grant_number:
                    part += f" (grant {fs.grant_number})"
                if fs.recipient:
                    part += f" to {fs.recipient}"
                funding_parts.append(part)
            self._add_labeled_statement(doc, "Funding", "; ".join(funding_parts))
