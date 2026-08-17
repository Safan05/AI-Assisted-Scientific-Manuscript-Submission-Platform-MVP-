# backend/app/services/parsing/docling_parser.py
"""
Document Parser Backend.

Provides structured parsing of scientific manuscripts (.docx) into a canonical
SectionNode hierarchy, extracting embedded figures and structured tables.

Supports:
1. Docling (IBM Research) when available.
2. High-fidelity python-docx fallback parser that extracts heading trees, body
   paragraphs, tables, and embedded images seamlessly without external dependencies.
"""

from __future__ import annotations

import io
import os
import re
import tempfile
import logging
from dataclasses import dataclass, field
from typing import Optional, Any
from pathlib import Path

try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import DocumentStream
except ImportError:
    DocumentConverter = None
    DocumentStream = None

try:
    import docx
    from docx.document import Document as _DocxDocument
except ImportError:
    docx = None

from app.schemas.manuscript_ir import SectionNode

logger = logging.getLogger(__name__)


@dataclass
class ExtractedPictureItem:
    """Represents a picture/figure extracted from the document."""
    index: int
    image_bytes: bytes
    mime_type: str = "image/png"
    caption: Optional[str] = None
    original_filename: Optional[str] = None
    bounding_box: Optional[dict[str, Any]] = None
    provenance: Optional[dict[str, Any]] = None


@dataclass
class DoclingParseResult:
    """Result container returned by DoclingParser."""
    sections: list[SectionNode] = field(default_factory=list)
    pictures: list[ExtractedPictureItem] = field(default_factory=list)
    tables: list[dict[str, Any]] = field(default_factory=list)
    full_markdown: str = ""
    raw_dict: dict[str, Any] = field(default_factory=dict)


class DoclingParser:
    """
    Parses scientific manuscripts into a structured SectionNode tree,
    extracting embedded figures and structured tables.
    """

    def __init__(self, converter: Optional[Any] = None):
        self.converter = None
        if converter is not None:
            self.converter = converter
        elif DocumentConverter is not None:
            try:
                self.converter = DocumentConverter()
            except Exception as e:
                logger.warning("Could not initialize Docling DocumentConverter: %s", e)
                self.converter = None

    def parse(self, file_bytes: bytes, filename: str = "manuscript.docx") -> DoclingParseResult:
        """
        Parse raw manuscript bytes into SectionNode tree, pictures, and tables.
        """
        # 1. Try Docling if available
        if self.converter is not None:
            try:
                return self._parse_with_docling(file_bytes, filename)
            except Exception as e:
                logger.warning("Docling conversion failed, falling back to python-docx parser: %s", e)

        # 2. Resilient fallback using python-docx
        return self._parse_with_docx(file_bytes, filename)

    def _parse_with_docling(self, file_bytes: bytes, filename: str) -> DoclingParseResult:
        suffix = Path(filename).suffix if Path(filename).suffix else ".docx"
        temp_file = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
        try:
            temp_file.write(file_bytes)
            temp_file.flush()
            temp_file.close()

            conversion_result = self.converter.convert(temp_file.name)
        finally:
            if os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                except OSError as e:
                    logger.warning("Failed to remove temporary file %s: %s", temp_file.name, e)

        doc = conversion_result.document

        try:
            full_markdown = doc.export_to_markdown()
        except Exception:
            full_markdown = ""

        try:
            raw_dict = doc.export_to_dict()
        except Exception:
            raw_dict = {}

        sections = self._build_section_tree_docling(doc)
        pictures = self._extract_pictures_docling(doc)
        tables = self._extract_tables_docling(doc)

        return DoclingParseResult(
            sections=sections,
            pictures=pictures,
            tables=tables,
            full_markdown=full_markdown,
            raw_dict=raw_dict,
        )

    def _build_section_tree_docling(self, doc: Any) -> list[SectionNode]:
        root_sections: list[SectionNode] = []
        stack: list[SectionNode] = []

        for item, level in doc.iterate_items():
            item_type = item.__class__.__name__
            label = str(getattr(item, "label", "")).lower()

            if "heading" in label or "header" in label or item_type == "SectionHeaderItem":
                heading_text = getattr(item, "text", "").strip()
                if not heading_text:
                    continue

                heading_level = getattr(item, "level", None)
                if heading_level is None or not isinstance(heading_level, int) or heading_level < 1:
                    heading_level = max(1, min(6, level if isinstance(level, int) else 1))

                node = SectionNode(
                    heading=heading_text,
                    level=heading_level,
                    content=[],
                    children=[],
                )

                while stack and stack[-1].level >= heading_level:
                    stack.pop()

                if stack:
                    stack[-1].children.append(node)
                else:
                    root_sections.append(node)

                stack.append(node)

            elif "paragraph" in label or "text" in label or item_type in ("TextItem", "ParagraphItem", "ListItem"):
                text_content = getattr(item, "text", "").strip()
                if not text_content:
                    continue

                if stack:
                    stack[-1].content.append(text_content)
                else:
                    if not root_sections or root_sections[0].heading != "__preamble__":
                        preamble_node = SectionNode(
                            heading="__preamble__",
                            level=0,
                            content=[],
                            children=[],
                        )
                        root_sections.insert(0, preamble_node)
                    root_sections[0].content.append(text_content)

            elif "table" in label or item_type == "TableItem":
                table_md = ""
                if hasattr(item, "export_to_markdown"):
                    try:
                        table_md = item.export_to_markdown()
                    except Exception:
                        table_md = ""

                caption = ""
                if hasattr(item, "caption") and item.caption:
                    caption = getattr(item.caption, "text", str(item.caption)).strip()

                table_entry = f"[Table: {caption}]\n{table_md}" if caption else table_md
                if table_entry.strip():
                    if stack:
                        stack[-1].content.append(table_entry)
                    else:
                        if not root_sections or root_sections[0].heading != "__preamble__":
                            root_sections.insert(0, SectionNode(heading="__preamble__", level=0, content=[], children=[]))
                        root_sections[0].content.append(table_entry)

        return root_sections

    def _extract_pictures_docling(self, doc: Any) -> list[ExtractedPictureItem]:
        pictures: list[ExtractedPictureItem] = []
        pic_index = 0

        for item, _ in doc.iterate_items():
            item_type = item.__class__.__name__
            label = str(getattr(item, "label", "")).lower()

            if "picture" in label or "image" in label or item_type == "PictureItem":
                pic_index += 1
                caption_text: Optional[str] = None
                if hasattr(item, "caption") and item.caption:
                    caption_text = getattr(item.caption, "text", str(item.caption)).strip()

                img_bytes: Optional[bytes] = None
                mime_type = "image/png"

                if hasattr(item, "image") and item.image:
                    try:
                        pil_image = item.image.pil_image if hasattr(item.image, "pil_image") else item.image
                        buf = io.BytesIO()
                        pil_image.save(buf, format="PNG")
                        img_bytes = buf.getvalue()
                    except Exception:
                        img_bytes = None

                if img_bytes:
                    pictures.append(
                        ExtractedPictureItem(
                            index=pic_index,
                            image_bytes=img_bytes,
                            mime_type=mime_type,
                            caption=caption_text,
                            original_filename=f"figure_{pic_index}.png",
                        )
                    )

        return pictures

    def _extract_tables_docling(self, doc: Any) -> list[dict[str, Any]]:
        tables: list[dict[str, Any]] = []
        table_idx = 0

        for item, _ in doc.iterate_items():
            item_type = item.__class__.__name__
            label = str(getattr(item, "label", "")).lower()

            if "table" in label or item_type == "TableItem":
                table_idx += 1
                caption_str = ""
                if hasattr(item, "caption") and item.caption:
                    caption_str = getattr(item.caption, "text", str(item.caption)).strip()

                markdown_repr = ""
                if hasattr(item, "export_to_markdown"):
                    try:
                        markdown_repr = item.export_to_markdown()
                    except Exception:
                        markdown_repr = ""

                tables.append({
                    "index": table_idx,
                    "caption": caption_str,
                    "markdown": markdown_repr,
                })

        return tables

    # ── High-Fidelity Python-DOCX Parser Fallback ───────────────────────────
    def _parse_with_docx(self, file_bytes: bytes, filename: str) -> DoclingParseResult:
        """
        Parses a .docx manuscript using python-docx, constructing structured sections,
        tables, embedded figures, and full markdown text.
        """
        if docx is None:
            raise RuntimeError("python-docx library is required to parse .docx documents.")

        doc = docx.Document(io.BytesIO(file_bytes))

        root_sections: list[SectionNode] = []
        stack: list[SectionNode] = []
        full_text_lines: list[str] = []

        # 1. Walk paragraphs and identify heading structure
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text_lines.append(text)
            style_name = para.style.name if para.style else ""

            # Check if this paragraph is a heading
            heading_level = None
            if style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    heading_level = 1
            elif style_name in ("Title", "Subtitle"):
                heading_level = 1
            elif re.match(r"^(?:[0-9]+\.|\b(?:[IVXLCDM]+)\.|\b(?:Introduction|Methods|Materials and Methods|Results|Discussion|Conclusion|Conclusions|Abstract|References|Acknowledgments|Funding|Conflict of interest|Author contributions)\b)", text, re.IGNORECASE) and len(text) < 120:
                heading_level = 1

            if heading_level is not None:
                heading_level = max(1, min(6, heading_level))
                node = SectionNode(
                    heading=text,
                    level=heading_level,
                    content=[],
                    children=[],
                )

                while stack and stack[-1].level >= heading_level:
                    stack.pop()

                if stack:
                    stack[-1].children.append(node)
                else:
                    root_sections.append(node)

                stack.append(node)
            else:
                if stack:
                    stack[-1].content.append(text)
                else:
                    if not root_sections or root_sections[0].heading != "__preamble__":
                        preamble_node = SectionNode(
                            heading="__preamble__",
                            level=0,
                            content=[],
                            children=[],
                        )
                        root_sections.insert(0, preamble_node)
                    root_sections[0].content.append(text)

        # 2. Extract structured tables
        tables_list: list[dict[str, Any]] = []
        for t_idx, table in enumerate(doc.tables, start=1):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_cells)

            # Build markdown table representation
            md_lines = []
            if table_rows:
                header = table_rows[0]
                md_lines.append("| " + " | ".join(header) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in table_rows[1:]:
                    md_lines.append("| " + " | ".join(row) + " |")
            table_md = "\n".join(md_lines)

            tables_list.append({
                "index": t_idx,
                "caption": f"Table {t_idx}",
                "markdown": table_md,
                "rows": table_rows,
            })

            # Append table markdown to current section or preamble
            if table_md:
                table_entry = f"[Table {t_idx}]\n{table_md}"
                if stack:
                    stack[-1].content.append(table_entry)
                elif root_sections:
                    root_sections[-1].content.append(table_entry)

        # 3. Extract embedded pictures / images from .docx parts
        pictures_list: list[ExtractedPictureItem] = []
        pic_idx = 0
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    img_bytes = image_part.blob
                    content_type = image_part.content_type or "image/png"
                    pic_idx += 1
                    pictures_list.append(
                        ExtractedPictureItem(
                            index=pic_idx,
                            image_bytes=img_bytes,
                            mime_type=content_type,
                            caption=f"Figure {pic_idx}",
                            original_filename=Path(rel.target_ref).name or f"figure_{pic_idx}.png",
                        )
                    )
        except Exception as e:
            logger.warning("Notice extracting images from docx parts: %s", e)

        full_markdown = "\n\n".join(full_text_lines)

        return DoclingParseResult(
            sections=root_sections,
            pictures=pictures_list,
            tables=tables_list,
            full_markdown=full_markdown,
            raw_dict={"source": "python-docx", "filename": filename},
        )
